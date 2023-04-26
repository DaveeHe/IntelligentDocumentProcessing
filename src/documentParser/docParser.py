# -*- coding: utf-8 -*-
# @Time : 2023/4/26 16:10
# @Author : Professor He
# @Email : hechangjia1018@163.com
# @File : docParser.py
# @Project : IntelligentDocumentProcessing

import os
import docx
import json
from bs4 import BeautifulSoup
from src.parser_document.paragraph_type_name import *
from src.parser_document.common_util import filter_page_number


def get_plain_text_content(filename):
    document = docx.Document(filename)
    paragraphs = document.paragraphs
    body_element = document._body._body
    p_lst = body_element.p_lst
    tables = document.tables
    articles = []
    num = len(paragraphs)
    for i in range(num):
        paragraph = paragraphs[i]
        line = {}
        text = paragraph.text
        text = text.strip()
        img_list = paragraph._element.xpath('.//pic:pic')
        ele = paragraph._p.getnext()
        if len(img_list) == 0 or not img_list:
            line['hasImage'] = False
        else:
            line['hasImage'] = True
        if ele.tag != '' and ele.tag[-3:] == 'tbl':
            line['hasTable'] = True
        else:
            line['hasTable'] = False
        paragraph_type = paragraph.style.name
        if paragraph_type is None or not text:
            line['paragraph_type'] = normal
        else:
            line['paragraph_type'] = paragraph_type
        line['style_id'] = paragraph.style.style_id
        if text or line['hasImage'] or line['hasTable']:
            line['text'] = text
            parts = paragraph.runs
            details = []
            for part in parts:
                detail = {}
                detail['bold'] = part.bold
                detail['text'] = part.text
                detail['font'] = part.font.name
                details.append(detail)
            line['detail'] = details
            articles.append(line)
        is_catalogue = True if paragraph_type and 'toc' in paragraph_type.lower() else False
        if is_catalogue:
            cur_element = p_lst[i]
            catalogue_txt = get_catalogue_content(cur_element)
            line['text'] = catalogue_txt
            details = []
            line['detail'] = details
            articles.append(line)
    return articles


def get_catalogue_content(element):
    soup = BeautifulSoup(element.xml, 'lxml')
    html_tags = soup.find_all(True)
    text = ''
    catalogue_mark = '............'
    for tag in html_tags:
        if tag.name == 'w:r':
            if tag.text.strip():
                txt = tag.text.strip()
                if 'toc' in txt.lower():
                    continue
                is_page = filter_page_number(txt)
                if text and is_page:
                    if catalogue_mark in text:
                        text = text.replace(catalogue_mark, '')
                    text += catalogue_mark
                text += txt
    return text


def get_head_type(paragraph_type):
    for head_type in name2head.keys():
        # if paragraph_type is None:
        #     break
        if head_type in paragraph_type:
            return head_type
    return


def get_normal_alias_type(paragraph_type):
    for alias in normalAlias:
        if alias in paragraph_type:
            return normal
    return


def get_document_paragraph_type(filename):
    articles = get_plain_text_content(filename)
    size = len(articles)
    new_articles = []
    for i in range(size):
        article = articles[i]
        paragraph_type = article['paragraph_type']
        if paragraph_type == '列出段落1' and article['style_id'] == '11':
            paragraph_type = normal
        elif paragraph_type == '样式2' and article['style_id'] == '27':
            paragraph_type = 'Heading 2'
        elif paragraph_type == '理文2' and (article['style_id'] == '21' or article['style_id'] == '41'):
            paragraph_type = 'Heading 2'
        elif paragraph_type == '13.5.11.1':
            paragraph_type = 'Heading 4'
        elif paragraph_type in name2Normal:
            paragraph_type = name2Normal[paragraph_type]
        elif paragraph_type in name2Caption:
            paragraph_type = name2Caption[paragraph_type]
        elif paragraph_type in name2title_catalogue:
            paragraph_type = name2title_catalogue[paragraph_type]
        elif paragraph_type in name2head:
            paragraph_type = name2head[paragraph_type]
        elif get_head_type(paragraph_type):
            paragraph_type = get_head_type(paragraph_type)
        elif get_normal_alias_type(paragraph_type):
            paragraph_type = get_normal_alias_type(paragraph_type)
        elif paragraph_type in appendix:
            pass
        elif paragraph_type in special_head:
            pass

        else:
            # paragraph_type = normal
            print(filename, paragraph_type)
            print(article['text'])
        line = {}
        line['text'] = article['text']
        line['hasTable'] = article['hasTable']
        line['hasImage'] = article['hasImage']
        line['paragraph_type'] = paragraph_type
        line['detail'] = article['detail']
        new_articles.append(line)

    return new_articles


def check_structure(articles):
    text_catalogue = set()
    text_head = set()
    new_article = []
    for line in articles:
        paragraph_type = line['paragraph_type']
        if paragraph_type in catalogue2head.keys():
            text_catalogue.add(paragraph_type)
            continue
        if paragraph_type in catalogue2head.values():
            text_head.add(paragraph_type)
            line['size'] = head2size[paragraph_type]
            new_article.append(line)
            continue
        line['size'] = -100
        new_article.append(line)
    for catalogue in text_catalogue:
        if catalogue2head[catalogue] not in text_head:
            print('the structure of the article is destroyed')
            return []

    return new_article


def get_article_title(articles, only_title=True):
    size = len(articles)
    titles = []
    for i in range(size):
        line = articles[i]
        if line['size'] > -100:
            titles.append((line, i))
    if only_title:
        return titles
    leaf_title = []
    size = len(titles)
    for i in range(size):
        cur = titles[i][0]
        if i < size - 1:
            next = titles[i + 1][0]
            if cur['size'] <= next['size']:
                leaf_title.append(titles[i])
            continue
        leaf_title.append(titles[i])
    return titles, leaf_title


def save_process_result(new_articles, file, directory):
    target = directory + file[:-4] + '.txt'
    with open(target, 'w', encoding='utf-8') as f:
        for line in new_articles:
            line = json.dumps(line, ensure_ascii=False)
            f.write(line)
            f.write('\n')


def parse_document_structure2file(filepath, file):
    articles = get_document_paragraph_type(filepath)
    articles_with_size = check_structure(articles)
    titles, leaf_title = get_article_title(articles_with_size)

    save_process_result(articles_with_size, file, './brochure_content/')
    save_process_result(titles, file, './brochure_title/')
    save_process_result(leaf_title, file, './brochure_leaf_title/')
    return titles, leaf_title


def get_pure_text_title_from_document(filepath):
    articles = get_document_paragraph_type(filepath)
    articles_with_size = check_structure(articles)
    titles = get_article_title(articles_with_size)
    pure_texts = [line['text'] for line in articles_with_size if line['text']]
    pure_titles = [paragraph[0]['text'] for paragraph in titles if paragraph[0]['text']]
    return pure_texts, pure_titles


def get_content_from_table(table):
    table_text = []
    for i in range(len(table.rows)):
        row_text = []
        for j in range(len(table.columns)):
            # print(f"{i}行{j}列：数据：{table.cell(i, j).text}")
            row_text.append(table.cell(i, j).text)
        table_text.append(row_text)


